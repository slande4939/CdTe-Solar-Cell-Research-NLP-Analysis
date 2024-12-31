#!/usr/bin/env python
# coding: utf-8

# In[1]:


# -*- coding: utf-8 -*-
"""
Spyder Editor

This is a temporary script file.
"""
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on 2/2/2024

@author:Sherly Lande
"""

###############################################################################
### packages required to run code.  Make sure to install all required packages.
###############################################################################
import re,string
from nltk.corpus import stopwords
from nltk.stem import PorterStemmer

from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from sklearn.metrics.pairwise import cosine_similarity

from sklearn.manifold import MDS


from sklearn.preprocessing import LabelEncoder, OneHotEncoder
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import roc_auc_score, accuracy_score, confusion_matrix
from sklearn.model_selection import cross_val_score, StratifiedKFold
from sklearn.tree import DecisionTreeClassifier
import matplotlib.pyplot as plt
from sklearn.model_selection import train_test_split, KFold
from multiprocessing import Pool, freeze_support
import pandas as pd
import os

from gensim.models import Word2Vec,LdaMulticore, TfidfModel
from gensim import corpora


from gensim.models.doc2vec import Doc2Vec, TaggedDocument


import numpy as np






# In[2]:


###############################################################################
### Function to process documents
###############################################################################
def clean_doc(doc): 
    #split document into individual words
    tokens=doc.split()
    re_punc = re.compile('[%s]' % re.escape(string.punctuation))
    # remove punctuation from each word
    tokens = [re_punc.sub('', w) for w in tokens]
    # remove remaining tokens that are not alphabetic
    tokens = [word for word in tokens if word.isalpha()]
    # filter out short tokens
    tokens = [word for word in tokens if len(word) > 4]
    #lowercase all words
    tokens = [word.lower() for word in tokens]
    # filter out stop words
    stop_words = set(stopwords.words('english'))
    tokens = [w for w in tokens if not w in stop_words]         
    # word stemming    
    # ps=PorterStemmer()
    # tokens=[ps.stem(word) for word in tokens]
    return tokens



# In[3]:


###############################################################################
#   Functions to label encoding
###############################################################################
def One_Hot(variable):
    LE=LabelEncoder()
    LE.fit(variable)
    Label1=LE.transform(variable)
    OHE=OneHotEncoder()
    labels=OHE.fit_transform(Label1.reshape(-1,1)).toarray()
    return labels, LE, OHE





# In[4]:


###############################################################################
### Processing text into lists
###############################################################################

#set working Directory to where class corpus is saved.
os.chdir(r'C:/Users/Lavieestbelle$1/Desktop/NLP-Project/Word document/')


#read in class corpus csv into python
data=pd.read_csv('Class_Corpus.csv')

#create empty list to store text documents titles
titles=[]

#for loop which appends the DSI title to the titles list
for i in range(0,len(data)):
    temp_text=data['DSI_Title'].iloc[i]
    titles.append(temp_text)

#create empty list to store text documents
text_body=[]

#for loop which appends the text to the text_body list
for i in range(0,len(data)):
    temp_text=data['Text'].iloc[i]
    text_body.append(temp_text)

#Note: the text_body is the unprocessed list of documents read directly form 
#the csv.
    
#empty list to store processed documents
processed_text=[]
#for loop to process the text to the processed_text list
for i in text_body:
    text=clean_doc(i)
    processed_text.append(text)

#Note: the processed_text is the PROCESSED list of documents read directly form 
#the csv.  Note the list of words is separated by commas.


#stitch back together individual words to reform body of text
final_processed_text=[]

for i in processed_text:
    temp_DSI=i[0]
    for k in range(1,len(i)):
        temp_DSI=temp_DSI+' '+i[k]
    final_processed_text.append(temp_DSI)
    
#Note: We stitched the processed text together so the TFIDF vectorizer can work.
#Final section of code has 3 lists used.  2 of which are used for further processing.
#(1) text_body - unused, (2) processed_text (used in W2V), 
#(3) final_processed_text (used in TFIDF), and (4) DSI titles (used in TFIDF Matrix)
 




# In[5]:


###############################################################################
### Sklearn TFIDF 
###############################################################################
#note the ngram_range will allow you to include multiple words within the TFIDF matrix
#Call Tfidf Vectorizer
Tfidf=TfidfVectorizer(ngram_range=(1,3))

#fit the vectorizer using final processed documents.  The vectorizer requires the 
#stiched back together document.

TFIDF_matrix=Tfidf.fit_transform(final_processed_text)     

#creating datafram from TFIDF Matrix
matrix=pd.DataFrame(TFIDF_matrix.toarray(), columns=Tfidf.get_feature_names_out(), index=titles)







# In[6]:


###############################################################################
### Explore TFIDF Values
###############################################################################

average_TFIDF={}
for i in matrix.columns:
    average_TFIDF[i]=np.mean(matrix[i])

average_TFIDF_DF=pd.DataFrame(average_TFIDF,index=[0]).transpose()

average_TFIDF_DF.columns=['TFIDF']

#calculate Q1 and Q3 range
Q1=np.percentile(average_TFIDF_DF, 25)
Q3=np.percentile(average_TFIDF_DF, 75)
IQR = Q3 - Q1
outlier=Q3+(1.5*IQR)


#words that exceed the Q3+IQR*1.5
outlier_list=average_TFIDF_DF[average_TFIDF_DF['TFIDF']>=outlier]



#can export matrix to csv and explore further if necessary


# In[7]:


###############################################################################
### Doc2Vec
###############################################################################
documents = [TaggedDocument(doc, [i]) for i, doc in enumerate(final_processed_text)]
model = Doc2Vec(documents, vector_size=100, window=2, min_count=1)

doc2vec_df=pd.DataFrame()
for i in range(0,len(processed_text)):
    vector=pd.DataFrame(model.infer_vector(processed_text[i])).transpose()
    doc2vec_df=pd.concat([doc2vec_df,vector], axis=0)

doc2vec_df=doc2vec_df.reset_index()

doc_titles={'title': titles}
t=pd.DataFrame(doc_titles)

doc2vec_df=pd.concat([doc2vec_df,t], axis=1)

doc2vec_df=doc2vec_df.drop('index', axis=1)



# In[8]:


###############################################################################
### Gensim Word2vec 
###############################################################################

#Note, there are opportunities to use the word2vec matrix to determine words 
#which are similar.  Similar words can be used to create equivalent classes.  
#k-means is not used to group individual words using the Word2Vec output.

#word to vec
model_w2v = Word2Vec(processed_text, vector_size=100, window=5, min_count=1)

#join all processed DSI words into single list
processed_text_w2v=[]
for i in processed_text:
    for k in i:
        processed_text_w2v.append(k)

#obtian all the unique words from DSI
w2v_words=list(set(processed_text_w2v))

#can also use the get_feature_names() from TFIDF to get the list of words
#w2v_words=Tfidf.get_feature_names()

#empty dictionary to store words with vectors
w2v_vectors={}

#for loop to obtain weights for each word
for i in w2v_words:
    temp_vec=model_w2v.wv[i]
    w2v_vectors[i]=temp_vec

#create a final dataframe to view word vectors
w2v_df=pd.DataFrame(w2v_vectors).transpose()


#the following section runs applies the k-means algorithm on the TFIDF matrix.


# In[9]:


###############################################################################
### K Means Clustering - TFIDF
###############################################################################
k=8
km = KMeans(n_clusters=k, random_state =89)
km.fit(TFIDF_matrix)
clusters = km.labels_.tolist()


terms = Tfidf.get_feature_names_out()
Dictionary={'Doc Name':titles, 'Cluster':clusters,  'Text': final_processed_text}
frame=pd.DataFrame(Dictionary, columns=['Cluster', 'Doc Name','Text'])

print("Top terms per cluster:")
#sort cluster centers by proximity to centroid
order_centroids = km.cluster_centers_.argsort()[:, ::-1] 

terms_dict=[]


#save the terms for each cluster and document to dictionaries.  To be used later
#for plotting output.

#dictionary to store terms and titles
cluster_terms={}
cluster_title={}


for i in range(k):
    print("Cluster %d:" % i),
    temp_terms=[]
    temp_titles=[]
    for ind in order_centroids[i, :10]:
        print(' %s' % terms[ind])
        terms_dict.append(terms[ind])
        temp_terms.append(terms[ind])
    cluster_terms[i]=temp_terms
    
    print("Cluster %d titles:" % i, end='')
    temp=frame[frame['Cluster']==i]
    for title in temp['Doc Name']:
        print(' %s,' % title, end='')
        temp_titles.append(title)
    cluster_title[i]=temp_titles





# In[10]:


###############################################################################
### Plotting
###############################################################################

# convert two components as we're plotting points in a two-dimensional plane
# "precomputed" because we provide a distance matrix
# we will also specify `random_state` so the plot is reproducible.


mds = MDS(n_components=2, dissimilarity="precomputed", random_state=1)

dist = 1 - cosine_similarity(TFIDF_matrix)

pos = mds.fit_transform(dist)  # shape (n_components, n_samples)

xs, ys = pos[:, 0], pos[:, 1]


#set up colors per clusters using a dict.  number of colors must correspond to K
cluster_colors = {0: 'black', 1: 'grey', 2: 'blue', 3: 'rosybrown', 4: 'firebrick', 
                  5:'red', 6:'darksalmon', 7:'sienna'}


#set up cluster names using a dict.  
cluster_dict=cluster_title

#create data frame that has the result of the MDS plus the cluster numbers and titles
df = pd.DataFrame(dict(x=xs, y=ys, label=clusters, title=range(0,len(clusters)))) 

#group by cluster
groups = df.groupby('label')

fig, ax = plt.subplots(figsize=(12, 12)) # set size
ax.margins(0.05) # Optional, just adds 5% padding to the autoscaling

#iterate through groups to layer the plot
#note that I use the cluster_name and cluster_color dicts with the 'name' lookup to return the appropriate color/label
for name, group in groups:
    ax.plot(group.x, group.y, marker='o', linestyle='', ms=12,
            label=cluster_dict[name], color=cluster_colors[name], 
            mec='none')
    ax.set_aspect('auto')
    ax.tick_params(\
        axis= 'x',          # changes apply to the x-axis
        which='both',      # both major and minor ticks are affected
        bottom='off',      # ticks along the bottom edge are off
        top='off',         # ticks along the top edge are off
        labelbottom='on')
    ax.tick_params(\
        axis= 'y',         # changes apply to the y-axis
        which='both',      # both major and minor ticks are affected
        left='off',      # ticks along the bottom edge are off
        top='off',         # ticks along the top edge are off
        labelleft='on')

ax.legend(loc='center left', bbox_to_anchor=(1, 0.5))      #show legend with only 1 point



# In[11]:


#The following section of code is to run the k-means algorithm on the doc2vec outputs.
#note the differences in document clusters compared to the TFIDF matrix.
###############################################################################
### K Means Clustering Doc2Vec
###############################################################################
doc2vec_k_means=doc2vec_df.drop('title', axis=1)

k=8
km = KMeans(n_clusters=k, random_state =89)
km.fit(doc2vec_k_means)
clusters_d2v = km.labels_.tolist()

Dictionary={'Doc Name':titles, 'Cluster':clusters_d2v,  'Text': final_processed_text}
frame=pd.DataFrame(Dictionary, columns=['Cluster', 'Doc Name','Text'])

#dictionary to store clusters and respective titles
cluster_title={}

#note doc2vec clusters will not have individual words due to the vector representation
#is based on the entire document not indvidual words. As a result, there won't be individual
#word outputs from each cluster.   
for i in range(k):
    temp=frame[frame['Cluster']==i]
    temp_title_list=[]
    for title in temp['Doc Name']:
        temp_title_list.append(title)
    cluster_title[i]=temp_title_list





# In[12]:


###############################################################################
### Plotting Doc2vec
###############################################################################
# convert two components as we're plotting points in a two-dimensional plane
# "precomputed" because we provide a distance matrix
# we will also specify `random_state` so the plot is reproducible.


mds = MDS(n_components=2, dissimilarity="precomputed", random_state=1)

dist = 1 - cosine_similarity(doc2vec_k_means)

pos = mds.fit_transform(dist)  # shape (n_components, n_samples)

xs, ys = pos[:, 0], pos[:, 1]


#set up colors per clusters using a dict.  number of colors must correspond to K
cluster_colors = {0: 'black', 1: 'grey', 2: 'blue', 3: 'rosybrown', 4: 'firebrick', 
                  5:'red', 6:'darksalmon', 7:'sienna'}


#set up cluster names using a dict.  
cluster_dict=cluster_title         

#create data frame that has the result of the MDS plus the cluster numbers and titles
df = pd.DataFrame(dict(x=xs, y=ys, label=clusters, title=range(0,len(clusters)))) 

#group by cluster
groups = df.groupby('label')

fig, ax = plt.subplots(figsize=(12, 12)) # set size
ax.margins(0.05) # Optional, just adds 5% padding to the autoscaling

#iterate through groups to layer the plot
#note that I use the cluster_name and cluster_color dicts with the 'name' lookup to return the appropriate color/label
for name, group in groups:
    ax.plot(group.x, group.y, marker='o', linestyle='', ms=12,
            label=cluster_dict[name], color=cluster_colors[name], 
            mec='none')
    ax.set_aspect('auto')
    ax.tick_params(\
        axis= 'x',          # changes apply to the x-axis
        which='both',      # both major and minor ticks are affected
        bottom='off',      # ticks along the bottom edge are off
        top='off',         # ticks along the top edge are off
        labelbottom='on')
    ax.tick_params(\
        axis= 'y',         # changes apply to the y-axis
        which='both',      # both major and minor ticks are affected
        left='off',      # ticks along the bottom edge are off
        top='off',         # ticks along the top edge are off
        labelleft='on')

ax.legend(loc='center left', bbox_to_anchor=(1, 0.5))      #show legend with only 1 point


#The following section is used to create a model to predict the clusters labels 
#based on the the TFIDF matrix and the doc2vec vectors.  Note the model performance 
#using the two different vectorization methods.


# In[13]:


###############################################################################
### Classification using various RF Model
###############################################################################
model_RF=RandomForestClassifier()


#TFIDF
Y=clusters
X=TFIDF_matrix

#cross validation
cv_score=cross_val_score(model_RF, X,Y, cv=2)

#mean CV score
np.mean(cv_score)


#Doc2Vec
Y=clusters_d2v
X=doc2vec_k_means

#cross validation
cv_score=cross_val_score(model_RF, X,Y, cv=2)

#mean CV score
np.mean(cv_score)


#the following section is example code to create ECs within the corpus.  A dictionary
#will need to be created for every EC.  Each EC will need to be applied to the corpus.
#Below is an example of how the function works.




# In[14]:


###############################################################################
### EC clean up code 
###############################################################################
def create_ec(dictionary, corpus):
    for key, values in dictionary.items():
        for value in values:
            corpus= corpus.replace(value, key)
    return corpus


corpus='i like swiss.  i like cheddar.  i like provolone.'
cheese_dic={'cheese': ['swiss', 'cheddar', 'provolone']}

corpus_new=create_ec(cheese_dic, corpus)





# In[15]:


###############################################################################
###  LDA Code
###############################################################################

#LDA using bag of words
dictionary = corpora.Dictionary(processed_text)
corpus = [dictionary.doc2bow(doc) for doc in processed_text]

ldamodel = LdaMulticore(corpus, num_topics=3, id2word=dictionary, passes=2)    

for idx, topic in ldamodel.print_topics(-1):
    print('Topic: {} \nWords: {}'.format(idx, topic))


#LDA using TFIDF
tfidf = TfidfModel(corpus)
corpus_tfidf = tfidf[corpus]
ldamodel = LdaMulticore(corpus_tfidf, num_topics=3, id2word=dictionary, passes=2)    

for idx, topic in ldamodel.print_topics(-1):
    print('Topic: {} \nWords: {}'.format(idx, topic))

   # Code to generate figures and charts
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.manifold import TSNE
from sklearn.decomposition import PCA

# Mock data generation for visualization purposes
# TF-IDF Matrix
np.random.seed(0)
tfidf_data = np.random.rand(10, 15)  # Simulating TF-IDF for 10 documents and 15 terms
tfidf_df = pd.DataFrame(tfidf_data, columns=[f"Term_{i}" for i in range(1, 16)])
documents = [f"Document_{i}" for i in range(1, 11)]

# Doc2Vec Vectors
doc2vec_data = np.random.rand(10, 100)  # Simulating Doc2Vec vectors for 10 documents
doc2vec_df = pd.DataFrame(doc2vec_data, index=documents)

# Word2Vec Vectors
word2vec_data = np.random.rand(15, 100)  # Simulating Word2Vec vectors for 15 terms
word2vec_df = pd.DataFrame(word2vec_data, index=[f"Term_{i}" for i in range(1, 16)])

# Visualization functions
def plot_tfidf_heatmap(df):
    plt.figure(figsize=(10, 8))
    sns.heatmap(df, annot=False, cmap="YlGnBu")
    plt.title("TF-IDF Matrix Heatmap")
    plt.xlabel("Terms")
    plt.ylabel("Documents")
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.show()

def plot_vectors_2d(df, title, labels=None):
    pca = PCA(n_components=2)
    result = pca.fit_transform(df)
    plt.figure(figsize=(8, 6))
    if labels is not None:
        plt.scatter(result[:, 0], result[:, 1], c=labels, cmap='viridis', marker='o')
    else:
        plt.scatter(result[:, 0], result[:, 1])
    plt.title(title)
    plt.xlabel("Component 1")
    plt.ylabel("Component 2")
    plt.grid(True)
    plt.tight_layout()
    plt.show()



# In[16]:


# Plotting TF-IDF Matrix Heatmap
plot_tfidf_heatmap(tfidf_df)


# In[17]:


# Plotting Doc2Vec Vectors in 2D
plot_vectors_2d(doc2vec_df, "Doc2Vec Vector Visualization")


# In[18]:


# Plotting Word2Vec Vectors in 2D
plot_vectors_2d(word2vec_df, "Word2Vec Vector Visualization")


# In[19]:


from sklearn.metrics import roc_curve, auc
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report

# Simulating classification data
X, y = np.random.rand(100, 10), np.random.randint(0, 2, 100)
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=0)

# Training a Random Forest Classifier
model_rf = RandomForestClassifier(random_state=0)
model_rf.fit(X_train, y_train)
y_pred = model_rf.predict(X_test)
y_pred_proba = model_rf.predict_proba(X_test)[:, 1]

# ROC Curve
fpr, tpr, thresholds = roc_curve(y_test, y_pred_proba)
roc_auc = auc(fpr, tpr)



# In[20]:


# Classification Report
report = classification_report(y_test, y_pred, output_dict=True)
accuracy = report['accuracy']
precision = report['weighted avg']['precision']
recall = report['weighted avg']['recall']



# In[21]:


# Plotting ROC Curve
plt.figure(figsize=(8, 6))
plt.plot(fpr, tpr, color='darkorange', lw=2, label=f'ROC curve (area = {roc_auc:.2f})')
plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
plt.xlim([0.0, 1.0])
plt.ylim([0.0, 1.05])
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('Receiver Operating Characteristic (ROC) Curve')
plt.legend(loc="lower right")
plt.grid(True)
plt.tight_layout()
plt.show()

accuracy, precision, recall


# In[22]:


# Generating visualizations based on the clustering results from the output document

# Data simulation for visualization purposes
# Assume 3 clusters with top terms and document distribution
clusters = {
    'Cluster 1': {'Top Terms': ['solar', 'energy', 'power'], 'Document Count': 10},
    'Cluster 2': {'Top Terms': ['cell', 'efficiency', 'material'], 'Document Count': 15},
    'Cluster 3': {'Top Terms': ['panel', 'cost', 'installation'], 'Document Count': 5},
}


# In[23]:


# Generating Term Frequency Bar Charts for Each Cluster
def plot_term_frequency(cluster_data):
    fig, axes = plt.subplots(1, 3, figsize=(15, 5), sharey=True)
    fig.suptitle('Top Terms Frequency in Each Cluster')
    for ax, (cluster, data) in zip(axes, cluster_data.items()):
        terms, frequencies = data['Top Terms'], range(len(data['Top Terms']), 0, -1)
        ax.bar(terms, frequencies)
        ax.set_title(cluster)
        ax.set_xlabel('Terms')
        ax.set_ylabel('Frequency')
    plt.tight_layout(rect=[0, 0.03, 1, 0.95])
    plt.show()

plot_term_frequency(clusters)


# In[24]:


# Generating Document Distribution Pie Chart
labels = clusters.keys()
sizes = [data['Document Count'] for data in clusters.values()]
plt.figure(figsize=(8, 6))
plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
plt.title('Document Distribution Across Clusters')
plt.show()


#                             
# ###################################### First Trial of Code after corpus creation ##########################################

# In[36]:


pip install docx2txt


# In[2]:


import numpy as np


# In[3]:


#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sunday, February 11, 2024

@author: Sherly Lande
@updated: pchan Feb 11, 2024
"""

import os
import docx2txt
import pandas as pd
import codecs
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from sklearn.decomposition import LatentDirichletAllocation, PCA
from sklearn.feature_extraction.text import CountVectorizer, TfidfVectorizer
from sklearn.cluster import KMeans
from gensim.models.doc2vec import Doc2Vec, TaggedDocument
import matplotlib.pyplot as plt

nltk.download('stopwords')
nltk.download('punkt')

# Function to remove stopwords
def remove_stopwords(text):
    stop_words = set(stopwords.words('english'))
    tokens = word_tokenize(text.lower())
    filtered_text = [word for word in tokens if word not in stop_words]
    return ' '.join(filtered_text)

# Set working directory
os.chdir(r'C:\Users\Lavieestbelle$1\Desktop\NLP-Project\Word document')
print("Working directory set.")

# Lists to store file name and body of text for txt files
file_name_txt = []
text_txt = []

for file in os.listdir('.'):
    if file.endswith('txt'):
        with codecs.open(file, 'r', encoding='utf-8', errors='ignore') as fdata:
            contents = fdata.read()
            cleaned_content = remove_stopwords(contents)
            text_txt.append(cleaned_content)
            file_name_txt.append(file)

print(f"Processed {len(file_name_txt)} txt files.")

# Create dictionary for txt corpus
corpus_txt = {'DSI_Title': file_name_txt, 'Text': text_txt}
c1 = pd.DataFrame(corpus_txt)

# Function to retrieve and turn document into text
def retrieve_DSI(file_name):
    text = docx2txt.process(file_name)
    return text

# Lists to store file name and body of text for docx files
file_name_docx = []
text_docx = []

# Loop to iterate through documents in working directory
for file in os.listdir('.'):
    if file.endswith('.docx'):
        text_body = retrieve_DSI(file)
        cleaned_text_body = remove_stopwords(text_body)
        file_name_docx.append(file)
        text_docx.append(cleaned_text_body)

print(f"Processed {len(file_name_docx)} docx files.")

# Create dictionary for docx corpus
corpus_docx = {'DSI_Title': file_name_docx, 'Text': text_docx}
c2 = pd.DataFrame(corpus_docx)

# Combine txt and docx corpora
c_final = pd.concat([c1, c2], axis=0)
print("Corpora combined.")

# Output a CSV containing the class corpus along with titles
c_final.to_csv('Class_Corpus.csv', index=True, encoding="utf-8-sig")
print("CSV file 'Class_Corpus.csv' created.")

# Read the CSV File
df = pd.read_csv("Class_Corpus.csv", encoding='utf-8')
documents = df['Text'].tolist()

# Apply LDA
count_vect = CountVectorizer(max_features=1000)
doc_term_matrix = count_vect.fit_transform(documents)

lda = LatentDirichletAllocation(n_components=5, random_state=0)
lda.fit(doc_term_matrix)
print("LDA model trained.")

# Apply Doc2Vec
tagged_documents = [TaggedDocument(words=word_tokenize(doc.lower()), tags=[str(i)]) for i, doc in enumerate(documents)]
d2v_model = Doc2Vec(vector_size=50, min_count=2, epochs=30)
d2v_model.build_vocab(tagged_documents)
d2v_model.train(tagged_documents, total_examples=d2v_model.corpus_count, epochs=d2v_model.epochs)
print("Doc2Vec model trained.")

# Printing the inferred vector for document 0
print(f"Inferred vector for document 0: {d2v_model.infer_vector(tagged_documents[0].words)}")

# Find Similar Documents to document 0
similar_documents = d2v_model.dv.most_similar(0)
print(f"Documents similar to document 0: {similar_documents}")

# Apply TF-IDF and K-Means Clustering
tfidf_vectorizer = TfidfVectorizer(max_features=1000)
tfidf_matrix = tfidf_vectorizer.fit_transform(documents)
print("TF-IDF matrix created.")

n_clusters = 5
kmeans = KMeans(n_clusters=n_clusters, random_state=0)
kmeans.fit(tfidf_matrix)
print("K-Means clustering applied.")

labels = kmeans.labels_

# Dimensionality Reduction
pca = PCA(n_components=2)
reduced_vectors = pca.fit_transform(tfidf_matrix.toarray())
print("Dimensionality reduction completed.")

# Plot
plt.figure(figsize=(10, 10))
for i in range(n_clusters):
    plt.scatter(reduced_vectors[labels == i, 0], reduced_vectors[labels == i, 1], label=f'Cluster {i}')
plt.legend()
plt.title("Document Clusters")
plt.xlabel("PCA 1")
plt.ylabel("PCA 2")
plt.show()
print("Clustering visualization completed.")

# Generating synthetic data for document lengths
document_lengths = np.random.lognormal(mean=2, sigma=0.75, size=1000)

plt.figure(figsize=(10, 6))
plt.hist(document_lengths, bins=30, color='skyblue', edgecolor='black')
plt.title('Document Length Distribution')
plt.xlabel('Document Length (number of words)')
plt.ylabel('Frequency')
plt.show()

# Synthetic data for word frequencies
words = ['data', 'analysis', 'model', 'learning', 'algorithm']
frequencies = [120, 90, 75, 60, 45]

plt.figure(figsize=(10, 6))
plt.bar(words, frequencies, color='lightgreen', edgecolor='black')
plt.title('Word Frequency Distribution')
plt.xlabel('Word')
plt.ylabel('Frequency')
plt.show()

# Synthetic data for topic distributions across 5 documents
topic_distributions = np.random.dirichlet(alpha=[0.5]*5, size=5)

plt.figure(figsize=(10, 6))
for i, dist in enumerate(topic_distributions):
    plt.plot(dist, label=f'Document {i+1}', marker='o', linestyle='-')
plt.title('Topic Distribution Across Documents')
plt.xlabel('Topic')
plt.ylabel('Proportion')
plt.legend()
plt.show()



# In[4]:


# Printing the inferred vector for document 0
print(f"Inferred vector for document 0: {d2v_model.infer_vector(tagged_documents[0].words)}")


# In[6]:


import matplotlib.pyplot as plt

# Inferred vector components for Document 0
inferred_vector = [-0.46146053, -0.5217497, -0.2896505, 0.18440129, -0.31862706, -0.37078035,
                   0.46896544, 0.986448, -0.7843045, -0.09377539, 0.5946158, -0.23484497,
                   -0.01734057, -0.03233296, -0.06463833, 0.07287369, 0.7267909, 0.58715135,
                   -0.99119455, 0.16423546, -0.26097956, -0.05481321, 0.2788756, 0.13157234,
                   0.5800375, 0.5597388, -0.40991023, -0.42835143, -0.8257464, -0.5932526,
                   0.22788079, 0.78099203, -0.2344651, 0.10444336, -0.8381515, 0.98092896,
                   -0.07317311, -0.82629555, 0.13319033, -0.53922343, 0.37636182, 0.37838796,
                   -0.04658436, -0.4085584, 0.69453585, 0.3454973, -0.2413884, -0.0415732,
                   0.48851112, 0.05921476]

# Plotting the bar chart
plt.figure(figsize=(12, 6))
plt.bar(range(len(inferred_vector)), inferred_vector)
plt.xlabel('Vector Components')
plt.ylabel('Component Magnitude')
plt.title('Inferred Vector for Document 0')
plt.grid(True)
plt.tight_layout()  # Adjust the layout to make room for the xlabel which may be cut off in some environments
plt.show()


# In[5]:


# Find Similar Documents to document 0
similar_documents = d2v_model.dv.most_similar(0)
print(f"Documents similar to document 0: {similar_documents}")


# In[7]:


# Assuming d2v_model is your trained Doc2Vec model and tagged_documents is your list of tagged documents
similar_documents = d2v_model.dv.most_similar(0)
print(f"Documents similar to document 0: {similar_documents}")

# The output should be the list of similar documents with their similarity scores


# In[27]:


# Plotting Doc2Vec Vectors in 2D
plot_vectors_2d(doc2vec_df, "Doc2Vec Vector Visualization")


# In[28]:


# Plotting Word2Vec Vectors in 2D
plot_vectors_2d(word2vec_df, "Word2Vec Vector Visualization")


# In[29]:


from sklearn.cluster import KMeans

# Simulating clustering with KMeans for TF-IDF and Doc2Vec
kmeans_tfidf = KMeans(n_clusters=3, random_state=0).fit(tfidf_data)
kmeans_doc2vec = KMeans(n_clusters=3, random_state=0).fit(doc2vec_data)

# Plotting KMeans Clustering Results for TF-IDF Matrix
plot_vectors_2d(tfidf_df, "TF-IDF Clustering Results", labels=kmeans_tfidf.labels_)

# Plotting KMeans Clustering Results for Doc2Vec Vectors
plot_vectors_2d(doc2vec_df, "Doc2Vec Clustering Results", labels=kmeans_doc2vec.labels_)


# In[30]:


from sklearn.metrics import roc_curve, auc
from sklearn.model_selection import train_test_split
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import classification_report

# Simulating classification data
X, y = np.random.rand(100, 10), np.random.randint(0, 2, 100)
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.3, random_state=0)

# Training a Random Forest Classifier
model_rf = RandomForestClassifier(random_state=0)
model_rf.fit(X_train, y_train)
y_pred = model_rf.predict(X_test)
y_pred_proba = model_rf.predict_proba(X_test)[:, 1]

# ROC Curve
fpr, tpr, thresholds = roc_curve(y_test, y_pred_proba)
roc_auc = auc(fpr, tpr)

# Classification Report
report = classification_report(y_test, y_pred, output_dict=True)
accuracy = report['accuracy']
precision = report['weighted avg']['precision']
recall = report['weighted avg']['recall']

# Plotting ROC Curve
plt.figure(figsize=(8, 6))
plt.plot(fpr, tpr, color='darkorange', lw=2, label=f'ROC curve (area = {roc_auc:.2f})')
plt.plot([0, 1], [0, 1], color='navy', lw=2, linestyle='--')
plt.xlim([0.0, 1.0])
plt.ylim([0.0, 1.05])
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('Receiver Operating Characteristic (ROC) Curve')
plt.legend(loc="lower right")
plt.grid(True)
plt.tight_layout()
plt.show()

accuracy, precision, recall


# In[31]:


import matplotlib.pyplot as plt

# Data simulation for visualization purposes
# Assume 3 clusters with top terms and document distribution
clusters = {
    'Cluster 1': {'Top Terms': ['solar', 'energy', 'power'], 'Document Count': 10},
    'Cluster 2': {'Top Terms': ['cell', 'efficiency', 'material'], 'Document Count': 15},
    'Cluster 3': {'Top Terms': ['panel', 'cost', 'installation'], 'Document Count': 5},
}

# Generating Term Frequency Bar Charts for Each Cluster
def plot_term_frequency(cluster_data):
    fig, axes = plt.subplots(1, 3, figsize=(15, 5), sharey=True)
    fig.suptitle('Top Terms Frequency in Each Cluster')
    for ax, (cluster, data) in zip(axes, cluster_data.items()):
        terms, frequencies = data['Top Terms'], [3, 2, 1]  # Assuming a simple descending frequency
        ax.bar(terms, frequencies)
        ax.set_title(cluster)
        ax.set_xlabel('Term')
        ax.set_ylabel('Frequency')
        ax.set_ylim(0, max(frequencies) + 1)  # Setting y limit for better visualization
    plt.tight_layout(rect=[0, 0.03, 1, 0.95])
    plt.show()

plot_term_frequency(clusters)

# Generating Document Distribution Pie Chart
labels = clusters.keys()
sizes = [data['Document Count'] for data in clusters.values()]
plt.figure(figsize=(8, 6))
plt.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=140)
plt.axis('equal')  # Equal aspect ratio ensures that pie is drawn as a circle.
plt.title('Document Distribution Across Clusters')
plt.show()


# In[32]:


import matplotlib.pyplot as plt
import numpy as np

# Sample data from the provided code
processed_files = 26  # Number of processed files
lda_topics = 5  # Number of topics in LDA
doc2vec_vector_size = 50  # Vector size for Doc2Vec
kmeans_clusters = 5  # Number of clusters in KMeans

# Data for a bar chart of processed files and model parameters
categories = ['Processed Files', 'LDA Topics', 'Doc2Vec Vector Size', 'KMeans Clusters']
values = [processed_files, lda_topics, doc2vec_vector_size, kmeans_clusters]

plt.figure(figsize=(10, 6))
plt.bar(categories, values, color=['blue', 'green', 'red', 'purple'])
plt.title('Summary of Processed Data and Model Parameters')
plt.ylabel('Count / Size')
for i, v in enumerate(values):
    plt.text(i, v + 0.5, str(v), ha='center', va='bottom')
plt.show()

# Assuming the LDA model has produced a topic distribution for a specific document
# Sample topic distribution for a document
lda_topic_distribution = np.random.rand(lda_topics)  # Randomly generated for illustration
lda_topic_distribution /= lda_topic_distribution.sum()  # Normalize to sum to 1

plt.figure(figsize=(10, 6))
plt.bar(range(lda_topics), lda_topic_distribution, color='orange')
plt.title('LDA Topic Distribution for a Sample Document')
plt.xlabel('Topic')
plt.ylabel('Distribution')
plt.xticks(range(lda_topics), [f'Topic {i+1}' for i in range(lda_topics)])
plt.show()



# In[33]:


from docx import Document

# Create a new Document
doc = Document()

# Add a title to the document
doc.add_heading('Ontology Structure for CdTe Solar Cell Technologies', 0)

# Data for the table
data = {
    "Fabrication": ["Between-film layers", "Film creation", "Doping strategies", "Material processing and integration"],
    "Fault Detection": ["Defect mechanism analysis", "Techniques for identifying faults"],
    "Quality Improvement": ["Efficiency enhancement techniques", "Material quality advancements"],
    "Measurements": ["Efficiency analysis", "Performance parameters"],
    "Life-cycle/Life-span": ["Environmental impact analysis", "Sustainability assessments"],
    "Technological Advancements": ["Carrier dynamics", "Threshold switching phenomena", "Buffer layer impacts"],
    "Environmental and Economic Aspects": ["Resource constraints", "Cost implications and market viability"],
    "Innovative Applications": ["Architectural applications (PV glazing)", "Integrated solar solutions"]
}

# Add the table to the document
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'

# Add header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Main Category'
hdr_cells[1].text = 'Specific Topics'

# Add the rest of the data to the table
for main_category, specific_topics in data.items():
    row_cells = table.add_row().cells
    row_cells[0].text = main_category
    row_cells[1].text = ", ".join(specific_topics)

# Save the document
doc_path = 'C:/Users/Lavieestbelle$1/Documents/Ontology_Structure_CdTe_Solar_Cell.docx'
doc.save(doc_path)


# In[ ]:




